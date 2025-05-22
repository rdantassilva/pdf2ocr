"""Functions for converting between different document formats."""

import io
import os
import subprocess
import tempfile
import time
from concurrent import futures
from typing import List

from docx import Document
from pdf2image import convert_from_path
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from tqdm import tqdm

from pdf2ocr.config import ProcessingConfig
from pdf2ocr.logging_config import log_message, setup_logging
from pdf2ocr.ocr import process_pdf_with_ocr
from pdf2ocr.state import shutdown_requested
from pdf2ocr.utils import timing_context


def _process_paragraphs(text):
    """Process and clean paragraphs from text input.

    Args:
        text (Union[str, List[str]]): Input text, can be a single string or list of strings

    Returns:
        List[str]: List of cleaned paragraphs
    """
    # If input is a list, join with double newlines
    if isinstance(text, list):
        text = "\n\n".join(text)

    # Split into paragraphs and process each one
    paragraphs = []
    for para in text.strip().split("\n\n"):
        if not para.strip():
            continue
        # Join lines with spaces and clean extra whitespace
        lines = [line.strip() for line in para.split("\n")]
        clean_text = " ".join(line for line in lines if line)
        if clean_text:
            paragraphs.append(clean_text)
    
    return paragraphs


def save_as_docx(text: List[str], output_path: str) -> float:
    """Saves extracted text to a DOCX file.

    Args:
        text: The text content to save. Can be a single string or a list of page texts.
        output_path: Path where to save the DOCX file

    Returns:
        float: Time taken to save the file in seconds
    """
    start = time.perf_counter()
    title = os.path.splitext(os.path.basename(output_path))[0].replace("_", " ")
    document = Document()
    document.core_properties.title = title
    document.core_properties.author = "pdf2ocr"

    # Add title as heading for multi-page documents
    if isinstance(text, list):
        document.add_heading(title, level=1)

    # Process and add paragraphs
    for para in _process_paragraphs(text):
        document.add_paragraph(para)

    document.save(output_path)
    return time.perf_counter() - start


def save_as_pdf(text_pages, output_path):
    """Creates a new PDF with OCR-extracted text in a clean, standardized format.

    This function generates a new PDF document that focuses on text readability
    and consistent formatting. It's ideal for cases where you want to extract
    and reformat content from PDFs (e.g., books, articles, documents).

    Key features:
    - Creates a new PDF with clean, consistent formatting
    - Adds simple page numbers for navigation
    - Uses standard fonts and margins for optimal readability
    - Creates output in 'pdf_ocr' directory
    - Includes PDF metadata (title, author, creation info)

    Technical approach:
    1. Uses reportlab to generate a new PDF from scratch
    2. Applies consistent formatting (margins, fonts, spacing)
    3. Preserves paragraph structure from OCR text
    4. Handles text flow and automatic pagination
    5. Adds document metadata

    Args:
        text_pages (list): List of text content for each page
        output_path (str): Path where to save the PDF file

    Returns:
        float: Time taken to save the file in seconds

    Note:
        This differs from process_layout_pdf_only() which preserves original layout.
        Choose this method when you want a clean, reformatted version of the document.
    """
    start = time.perf_counter()

    # Get title from filename
    title = os.path.splitext(os.path.basename(output_path))[0].replace("_", " ")

    # Create PDF with metadata
    c = canvas.Canvas(output_path, pagesize=A4)
    c.setTitle(title)
    c.setAuthor("pdf2ocr")
    c.setCreator("pdf2ocr")
    c.setSubject("OCR processed document")
    c.setKeywords("OCR, PDF, text recognition")

    width, height = A4

    # Skip empty pages at the start
    page_offset = 0
    for idx, page_text in enumerate(text_pages):
        if page_text.strip():
            break
        page_offset += 1

    for idx, page_text in enumerate(text_pages[page_offset:], start=1):
        x = 2 * cm  # Left margin
        y = height - 3 * cm  # Top margin

        # Page header
        c.setFont("Helvetica", 10)  # Changed to regular weight for consistency
        header = f"pdf2ocr - Page {idx}"  # New header format
        c.drawString(x, height - 1 * cm, header)  # Aligned left like HTML

        # Main text
        c.setFont("Helvetica", 10)
        line_height = 12

        # Process paragraphs
        for para in _process_paragraphs(page_text):
            # Split long paragraphs into lines that fit the page width
            words = para.split()
            current_line = []
            
            for word in words:
                current_line.append(word)
                line = " ".join(current_line)
                
                # Check if line width exceeds page width
                if c.stringWidth(line, "Helvetica", 10) > (width - 4 * cm):
                    # Remove last word and draw current line
                    current_line.pop()
                    line = " ".join(current_line)
                    
                    # Draw line and check page break
                    if y < 2 * cm:  # Check page end
                        c.showPage()
                        y = height - 3 * cm
                        # Repeat header on new page
                        c.setFont("Helvetica", 10)
                        c.drawString(x, height - 1 * cm, header)
                        c.setFont("Helvetica", 10)
                    
                    c.drawString(x, y, line)
                    y -= line_height
                    
                    # Start new line with the word that didn't fit
                    current_line = [word]
            
            # Draw remaining words in the last line
            if current_line:
                line = " ".join(current_line)
                if y < 2 * cm:  # Check page end
                    c.showPage()
                    y = height - 3 * cm
                    # Repeat header on new page
                    c.setFont("Helvetica", 10)
                    c.drawString(x, height - 1 * cm, header)
                    c.setFont("Helvetica", 10)
                
                c.drawString(x, y, line)
                y -= line_height
                page_has_content = True
            
            # Extra space between paragraphs
            y -= line_height

        if page_has_content:
            c.showPage()

    c.save()
    return time.perf_counter() - start


def save_as_html(text, output_path):
    """Saves extracted text to an HTML file.

    Args:
        text (Union[str, List[str]]): The text content to save. Can be a single string or a list of page texts.
        output_path (str): Path where to save the HTML file

    Returns:
        float: Time taken to save the file in seconds
    """
    start = time.perf_counter()
    try:
        title = os.path.splitext(os.path.basename(output_path))[0].replace("_", " ")

        # Start HTML content
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
        h1 {{ color: #333; }}
        .page {{ margin-bottom: 30px; }}
        .page-number {{ color: #666; font-style: italic; margin-bottom: 10px; }}
        p {{ margin-bottom: 15px; }}
    </style>
</head>
<body>"""

        # Handle multi-page documents
        if isinstance(text, list):
            # Add title
            html_content += f"\n<h1>{title}</h1>\n"

            # Skip empty pages at the start
            page_offset = 0
            for idx, page_text in enumerate(text):
                if page_text.strip():
                    break
                page_offset += 1

            # Process each page
            for page_num, page_text in enumerate(text[page_offset:], start=1):
                html_content += '\n<div class="page">\n'
                html_content += f'<div class="page-number">Page {page_num}</div>\n'

                # Process paragraphs in this page
                for para in _process_paragraphs(page_text):
                    html_content += f"<p>{para}</p>\n"

                html_content += "</div>\n"
        else:
            # Process single text string
            for para in _process_paragraphs(text):
                html_content += f"<p>{para}</p>\n"

        # Close HTML
        html_content += "</body>\n</html>"

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

    except Exception as e:
        raise Exception(f"Failed to save HTML: {e}")

    return time.perf_counter() - start


def convert_docx_to_epub(docx_path, epub_path, lang):
    """Converts DOCX to EPUB using Calibre.

    Args:
        docx_path (str): Path to source DOCX file
        epub_path (str): Path where to save the EPUB file
        lang (str): Language code for the document

    Returns:
        tuple: (success: bool, time_taken: float, output: str)
    """
    start = time.perf_counter()
    try:
        title = os.path.splitext(os.path.basename(str(epub_path)))[0].replace("_", " ")

        # Expanded language mapping
        TESS_TO_CALIBRE_LANG = {
            "por": "pt-BR",  # Brazilian Portuguese
            "eng": "en-US",  # US English
            "spa": "es",  # Spanish
            "fra": "fr",  # French
            "deu": "de-DE",  # German
            "ita": "it",  # Italian
            "nld": "nl",  # Dutch
            "rus": "ru",  # Russian
            "tur": "tr",  # Turkish
            "jpn": "ja",  # Japanese
            "chi_sim": "zh-CN",  # Simplified Chinese
            "chi_tra": "zh-TW",  # Traditional Chinese
            "chi_sim_vert": "zh-CN",  # Simplified Chinese (vertical)
            "chi_tra_vert": "zh-TW",  # Traditional Chinese (vertical)
            "heb": "he",  # Hebrew
        }

        tess_lang = lang.lower()  # Normalize language code
        calibre_lang = TESS_TO_CALIBRE_LANG.get(tess_lang)

        # Build command with improved options
        cmd = [
            "ebook-convert",
            str(docx_path),  # Ensure path is string
            str(epub_path),  # Ensure path is string
            "--title",
            title,
            "--authors",
            "pdf2ocr",
            "--comments",
            "Converted by pdf2ocr",
            "--chapter",
            "//h:h1",  # Use h1 headings as chapter markers
            "--chapter-mark",
            "pagebreak",  # Add page breaks at chapters
            "--level1-toc",
            "//h:h1",  # Use h1 headings in TOC
            "--toc-threshold",
            "1",  # Include all headings in TOC
            "--max-toc-links",
            "50",  # Reasonable limit for TOC entries
            "--pretty-print",  # Format HTML for better readability
        ]

        if calibre_lang:
            cmd.extend(["--language", calibre_lang])

        # Check if source file exists
        if not os.path.exists(docx_path):
            return (
                False,
                time.perf_counter() - start,
                f"Source file not found: {docx_path}",
            )

        # Ensure output directory exists
        os.makedirs(os.path.dirname(epub_path), exist_ok=True)

        # Run Calibre command and capture output
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            check=True
        )

        # Check if output file exists
        if not os.path.exists(epub_path):
            return False, time.perf_counter() - start, "Output file not created"

        # Combine stdout and stderr for logging
        output = result.stdout
        if result.stderr:
            output = f"{output}\n{result.stderr}"

        return True, time.perf_counter() - start, output

    except subprocess.CalledProcessError as e:
        error_msg = f"EPUB conversion failed: {e.stdout}"
        if e.stderr:
            error_msg += f"\nError details: {e.stderr}"
        if "ebook-convert: command not found" in str(e):
            error_msg = "Calibre (ebook-convert) not found. Please install Calibre and ensure it's in your PATH."
        return False, time.perf_counter() - start, error_msg
    except Exception as e:
        return False, time.perf_counter() - start, f"Unexpected error: {str(e)}"


def process_single_layout_pdf(filename: str, config: ProcessingConfig) -> tuple:
    """Process a single PDF file in layout-preserving mode.

    Args:
        filename: Name of the PDF file to process
        config: Processing configuration

    Returns:
        tuple: (success: bool, processing_time: float, error_message: str, log_messages: list)
    """
    # Setup logging for this process
    logger = setup_logging(config.log_path, config.quiet, is_worker=True)
    log_messages = []

    try:
        pdf_path = os.path.join(config.source_dir, filename)
        base_name = os.path.splitext(filename)[0]
        out_path = os.path.join(config.pdf_dir, f"{base_name}_ocr.pdf")
        temp_pdf_path = os.path.join(config.pdf_dir, f"{base_name}_temp.pdf")

        with timing_context(f"Processing {filename}", logger) as get_file_time:
            # Convert PDF to images with high DPI for better OCR
            pages = convert_from_path(pdf_path, dpi=200)

            # Process each page with OCR
            pdf_pages = []

            # Create temporary directory for processing
            with tempfile.TemporaryDirectory() as temp_dir:
                with timing_context("OCR processing", None) as get_ocr_time:
                    for page_num, page_img in enumerate(
                        tqdm(pages, desc="Processing pages", unit="page")
                    ):
                        # Save image to temporary file with high quality for OCR
                        img_path = os.path.join(temp_dir, f"page_{page_num}.png")
                        page_img.save(img_path, "PNG")

                        # Generate PDF with OCR using tesseract
                        pdf_path = os.path.join(temp_dir, f"page_{page_num}")
                        cmd = [
                            "tesseract",
                            img_path,
                            pdf_path,
                            "-l",
                            config.lang,
                            "--dpi",
                            "200",
                            "pdf",
                        ]
                        subprocess.run(cmd, check=True, capture_output=True)

                        # Read generated PDF
                        with open(f"{pdf_path}.pdf", "rb") as f:
                            pdf_pages.append(f.read())

                log_messages = [
                    (
                        "INFO",
                        f"  OCR processing took {get_ocr_time.duration:.2f} seconds",
                    )
                ]

                # Merge PDF pages into temporary file
                with timing_context("PDF merging", None) as get_merge_time:
                    writer = PdfWriter()
                    writer.compress = True

                    for page_pdf in pdf_pages:
                        reader = PdfReader(io.BytesIO(page_pdf))
                        writer.add_page(reader.pages[0])

                    with open(temp_pdf_path, "wb") as fout:
                        writer.write(fout)

                # Compress the final PDF using Ghostscript
                with timing_context("PDF compression", None) as get_compress_time:
                    cmd = [
                        "gs",
                        "-sDEVICE=pdfwrite",
                        "-dCompatibilityLevel=1.4",
                        "-dPDFSETTINGS=/ebook",  # Balance between quality and size
                        "-dNOPAUSE",
                        "-dQUIET",
                        "-dBATCH",
                        f"-sOutputFile={out_path}",
                        temp_pdf_path
                    ]
                    subprocess.run(cmd, check=True, capture_output=True)
                    
                    # Remove temporary PDF
                    if os.path.exists(temp_pdf_path):
                        os.remove(temp_pdf_path)

                log_messages.append(
                    (
                        "INFO",
                        f"  Layout-preserving PDF created and compressed in {get_merge_time.duration + get_compress_time.duration:.2f} seconds",
                    )
                )

            return True, get_file_time.stop(), None, log_messages

    except Exception as e:
        error_msg = f"Error in {filename} during {e.__class__.__name__}: {str(e)}"
        log_messages.append(("ERROR", error_msg))
        # Clean up temporary file if it exists
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)
        return False, 0, error_msg, log_messages


def process_layout_pdf_only(config: ProcessingConfig, logger) -> None:
    """Generate searchable PDFs while preserving the original document layout.

    This function creates a searchable PDF that maintains the exact visual appearance
    of the original document. It's ideal for documents where layout and formatting
    are crucial (e.g., forms, scientific papers, magazines).

    Key features:
    - Preserves original document layout, fonts, and formatting
    - Adds invisible OCR text layer for searchability
    - Uses lower DPI (200) to optimize file size
    - Creates output in 'pdf_ocr_layout' directory
    """
    try:
        # Show warning about layout preservation mode
        log_message(
            logger,
            "WARNING",
            "Layout preservation mode only supports PDF output. Other formats will be disabled.",
            quiet=config.quiet,
        )

        # Create output directory
        os.makedirs(config.pdf_dir, exist_ok=True)
        log_message(
            logger,
            "DEBUG",
            f"Preserve-layout PDF folder created - {config.pdf_dir}",
            quiet=config.quiet,
        )
        log_message(
            logger, "DEBUG", "", quiet=config.quiet
        )  # Add empty line after directory creation

        # Get list of PDF files
        pdf_files = sorted(
            f for f in os.listdir(config.source_dir) if f.lower().endswith(".pdf")
        )
        if not pdf_files:
            log_message(logger, "WARNING", "No PDF files found!", quiet=config.quiet)
            return

        # Process files in parallel
        with timing_context("Total execution", logger) as get_total_time:
            # Use configured number of workers
            max_workers = config.workers
            log_message(
                logger,
                "INFO",
                f"Processing {len(pdf_files)} files using {max_workers} worker{'s' if max_workers > 1 else ''}",
                quiet=config.quiet,
            )
            log_message(
                logger, "INFO", "", quiet=config.quiet
            )  # Empty line for readability

            # Track processing statistics
            completed = 0
            successful = 0
            failed = 0
            errors = []

            from concurrent import futures
            with futures.ProcessPoolExecutor(max_workers=max_workers) as executor:
                # Submit all files for processing
                future_to_file = {
                    executor.submit(
                        process_single_layout_pdf, filename, config
                    ): filename
                    for filename in pdf_files
                }

                # Process results as they complete
                for future in futures.as_completed(future_to_file):
                    filename = future_to_file[future]
                    completed += 1

                    try:
                        success, processing_time, error, log_messages = future.result()

                        # Log file start
                        log_message(
                            logger,
                            "INFO",
                            f"\n[{completed}/{len(pdf_files)}] Processing: {filename}",
                            quiet=config.quiet,
                        )

                        # Write all log messages from the child process
                        for level, message in log_messages:
                            log_message(
                                logger, level, "  " + message, quiet=config.quiet
                            )

                        if success:
                            successful += 1
                            log_message(
                                logger,
                                "INFO",
                                f"  ✓ Completed successfully in {processing_time:.2f} seconds",
                                quiet=config.quiet,
                            )
                        else:
                            failed += 1
                            if error:
                                errors.append((filename, error))
                                log_message(
                                    logger,
                                    "ERROR",
                                    f"  ✗ Failed: {error}",
                                    quiet=config.quiet,
                                )

                        # Add empty line for readability
                        log_message(logger, "INFO", "", quiet=config.quiet)

                    except Exception as e:
                        failed += 1
                        error_msg = f"Error processing {filename}: {str(e)}"
                        log_message(logger, "ERROR", error_msg, quiet=config.quiet)
                        errors.append((filename, error_msg))

            # Log final summary
            total_time = get_total_time()
            summary = [
                "Processing Summary:",
                "----------------",
                f"Files processed: {completed}",
                f"Total time: {total_time:.2f} seconds",
            ]

            if completed > 0:
                avg_time = total_time / completed
                summary.append(f"Average time per file: {avg_time:.2f} seconds")

            summary.extend([f"Successful: {successful}", f"Failed: {failed}"])

            if errors:
                summary.extend(["Errors:", "-------"])
                for filename, error in errors:
                    summary.append(f"• {filename}:")
                    summary.append(f"  {error}")

            log_message(logger, "INFO", "\n".join(summary), quiet=config.quiet)

    except Exception as e:
        log_message(
            logger,
            "ERROR",
            f"Error processing layout PDF: {str(e)}",
            quiet=config.quiet,
        )
        raise


def process_single_pdf(filename: str, config: ProcessingConfig) -> tuple:
    """Process a single PDF file.

    Args:
        filename: Name of the PDF file to process
        config: Processing configuration

    Returns:
        tuple: (success: bool, processing_time: float, error_message: str, log_messages: list)
    """
    # Setup logging for this process
    logger = setup_logging(config.log_path, config.quiet, is_worker=True)
    log_messages = []

    try:
        pdf_path = os.path.join(config.source_dir, filename)
        base_name = os.path.splitext(filename)[0]

        with timing_context(f"Processing {filename}", logger) as get_file_time:
            try:
                # Process PDF pages with OCR
                with timing_context(
                    "OCR processing", None
                ) as get_ocr_time:  # Remove logger to avoid immediate logging
                    pages = process_pdf_with_ocr(
                        pdf_path,
                        config.lang,
                        logger,
                        config.quiet,
                        config.summary_output,
                    )
                    page_texts = [text for _, text in pages]

                # Generate output files based on configuration
                output_files = []

                if config.generate_pdf:
                    pdf_path = os.path.join(config.pdf_dir, base_name + "_ocr.pdf")
                    with timing_context("PDF generation", None) as get_pdf_time:
                        save_as_pdf(page_texts, pdf_path)
                        output_files.append(("PDF", get_pdf_time()))

                if config.generate_html:
                    html_path = os.path.join(config.html_dir, base_name + ".html")
                    with timing_context("HTML generation", None) as get_html_time:
                        save_as_html(page_texts, html_path)
                        output_files.append(("HTML", get_html_time()))

                if config.generate_docx:
                    docx_path = os.path.join(config.docx_dir, base_name + ".docx")
                    with timing_context("DOCX generation", None) as get_docx_time:
                        save_as_docx(page_texts, docx_path)
                        output_files.append(("DOCX", get_docx_time()))

                    if config.generate_epub:
                        epub_path = os.path.join(config.epub_dir, base_name + ".epub")
                        with timing_context("EPUB generation", None) as get_epub_time:
                            success, _, output = convert_docx_to_epub(docx_path, epub_path, config.lang)
                            # Log Calibre output only to file
                            if output and logger:
                                log_message(logger, "INFO", "Calibre output:", quiet=True)
                                log_message(logger, "INFO", output, quiet=True)
                            if success:
                                output_files.append(("EPUB", get_epub_time()))
                            else:
                                raise Exception(f"EPUB conversion failed: {output}")

                # Generate log messages for all successful outputs
                log_messages = [
                    (
                        "INFO",
                        f"  OCR processing took {get_ocr_time.duration:.2f} seconds",
                    )
                ]  # Add OCR timing first
                log_messages.extend(
                    [
                        ("INFO", f"  {format} created in {time:.2f} seconds")
                        for format, time in output_files
                    ]
                )

                return True, get_file_time.stop(), None, log_messages

            except Exception as e:
                # Add context to the error message
                error_msg = (
                    f"Error in {filename} during {e.__class__.__name__}: {str(e)}"
                )
                log_messages.append(("ERROR", error_msg))
                return False, get_file_time.stop(), error_msg, log_messages

    except Exception as e:
        # Handle setup/initialization errors
        error_msg = f"Failed to initialize processing for {filename}: {str(e)}"
        log_messages.append(("ERROR", error_msg))
        return False, 0, error_msg, log_messages

    finally:
        # Ensure logger is properly closed
        if logger and hasattr(logger, "close"):
            logger.close()


def process_pdfs_with_ocr(config: ProcessingConfig, logger):
    """Main function that orchestrates the entire conversion process"""

    # Helper function to determine if we should show a message in console
    def should_show_message():
        if config.quiet:  # In quiet mode, show nothing
            return False
        if config.summary_output:  # In summary mode, show nothing except summary
            return False
        return True  # In normal mode, show everything

    # Validate configuration
    config.validate()

    try:
        # Create output directories
        os.makedirs(config.get_effective_dest_dir(), exist_ok=True)
        if config.generate_docx:
            os.makedirs(config.docx_dir, exist_ok=True)
            log_message(
                logger,
                "DEBUG",
                f"DOCX folder created - {config.docx_dir}",
                quiet=not should_show_message(),
            )
        if config.generate_pdf:
            os.makedirs(config.pdf_dir, exist_ok=True)
            log_message(
                logger,
                "DEBUG",
                f"PDF folder created - {config.pdf_dir}",
                quiet=not should_show_message(),
            )
        if config.generate_epub:
            os.makedirs(config.epub_dir, exist_ok=True)
            log_message(
                logger,
                "DEBUG",
                f"EPUB folder created - {config.epub_dir}",
                quiet=not should_show_message(),
            )
        if config.generate_html:
            os.makedirs(config.html_dir, exist_ok=True)
            log_message(
                logger,
                "DEBUG",
                f"HTML folder created - {config.html_dir}",
                quiet=not should_show_message(),
            )

        # Add empty line after folder creation
        log_message(logger, "DEBUG", "", quiet=not should_show_message())

        # Get list of PDF files
        pdf_files = sorted(
            f for f in os.listdir(config.source_dir) if f.lower().endswith(".pdf")
        )
        if not pdf_files:
            log_message(
                logger, "WARNING", "No PDF files found!", quiet=False
            )  # Always show warnings
            return

        # Process files in parallel
        with timing_context("Total execution", logger) as get_total_time:
            # Use configured number of workers
            max_workers = config.workers
            log_message(
                logger,
                "INFO",
                f"Processing {len(pdf_files)} files using {max_workers} worker{'s' if max_workers > 1 else ''}",
                quiet=not should_show_message(),
            )
            log_message(
                logger, "INFO", "", quiet=not should_show_message()
            )  # Empty line for readability

            # Track processing statistics
            completed = 0
            successful = 0
            failed = 0
            errors = []

            with futures.ProcessPoolExecutor(max_workers=max_workers) as executor:
                # Submit all files for processing
                future_to_file = {
                    executor.submit(process_single_pdf, filename, config): filename
                    for filename in pdf_files
                }

                # Process results as they complete
                try:
                    for future in futures.as_completed(future_to_file):
                        if shutdown_requested:
                            executor.shutdown(wait=False)
                            break

                        filename = future_to_file[future]
                        completed += 1

                        try:
                            success, processing_time, error, log_messages = (
                                future.result()
                            )

                            # Log file start and messages
                            log_message(
                                logger,
                                "INFO",
                                f"\n[{completed}/{len(pdf_files)}] Processing: {filename}",
                                quiet=not should_show_message(),
                            )

                            # Write all log messages from the child process
                            for level, message in log_messages:
                                log_message(
                                    logger,
                                    level,
                                    "  " + message,
                                    quiet=not should_show_message(),
                                )

                            if success:
                                successful += 1
                                log_message(
                                    logger,
                                    "INFO",
                                    f"  ✓ Completed successfully in {processing_time:.2f} seconds",
                                    quiet=not should_show_message(),
                                )
                            else:
                                failed += 1
                                if error:
                                    errors.append((filename, error))
                                    log_message(
                                        logger,
                                        "ERROR",
                                        f"  ✗ Failed: {error}",
                                        quiet=False,
                                    )  # Always show errors

                            # Add empty line for readability
                            log_message(
                                logger, "INFO", "", quiet=not should_show_message()
                            )

                        except Exception as e:
                            failed += 1
                            error_msg = f"Error processing {filename}: {str(e)}"
                            log_message(
                                logger, "ERROR", error_msg, quiet=False
                            )  # Always show errors
                            errors.append((filename, error_msg))

                except KeyboardInterrupt:
                    executor.shutdown(wait=False)
                    log_message(
                        logger,
                        "WARNING",
                        "\nProcessing interrupted by user.",
                        quiet=False,
                    )  # Always show interruption
                    return

            if shutdown_requested:
                log_message(
                    logger, "WARNING", "\nProcessing interrupted by user.", quiet=False
                )  # Always show interruption
                return

            # Log final summary
            total_time = get_total_time()
            summary = [
                "Processing Summary:",
                "----------------",
                f"Files processed: {completed}",
                f"Total time: {total_time:.2f} seconds",
            ]

            if completed > 0:
                avg_time = total_time / completed
                summary.append(f"Average time per file: {avg_time:.2f} seconds")

            summary.extend([f"Successful: {successful}", f"Failed: {failed}"])

            if errors:
                summary.extend(["Errors:", "-------"])
                for filename, error in errors:
                    summary.append(f"• {filename}:")
                    summary.append(f"  {error}")

            log_message(logger, "INFO", "\n".join(summary), quiet=config.quiet)

    except Exception as e:
        log_message(
            logger, "ERROR", f"Error processing PDFs: {str(e)}", quiet=False
        )  # Always show errors
        raise
