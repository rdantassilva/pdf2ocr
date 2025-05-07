#coding:utf-8

"""
Script for OCR on PDFs generating DOCX, OCR-processed PDF and EPUB.

Description of Libraries:
- 📜 PIL/Pillow: Image processing (open, convert and enhance images for OCR)
- 🔍 pytesseract: Python interface for Tesseract OCR (text recognition in images)
- 📊 tqdm: Displays progress bars for long operations
- 📄 pdf2image: Converts PDF pages to images for processing
- 📝 python-docx: Creates and manipulates Word documents (.docx)
- 🖨️ reportlab: Programmatic PDF generation
- 🏷️ tesseract-ocr: OCR engine (optical character recognition)
- 🖼️ poppler-utils: PDF manipulation tools (includes pdftoppm)
- 📚 calibre: E-book management suite (provides ebook-convert)
- 🐍 PyPDF2: Merge per-page searchable PDFs into one (layout-preserving PDF mode)

Conversion Flow:
1. 📄 pdf2image uses `pdftoppm` to convert PDFs to images
2. 🖼️ pytesseract uses `tesseract-ocr` to extract text from images
3. 📝 python-docx generates Word documents with extracted text
4. 🖨️ reportlab creates PDFs with OCR-processed text
5. 📚 calibre (optional) converts DOCX to EPUB using `ebook-convert`

Important Notes:
- To generate EPUB, Calibre must be installed on the system
- The `ebook-convert` command must be available in PATH
- Very large files may require more time and memory
- OCR quality depends on the original PDF clarity
- The script creates separate folders for each output type (docx, pdf, epub)
"""

import os          # Interact with the file system (paths, folders, environment)
import io          # Handle in-memory binary streams for PDF merging (Used in --preserve-mode)
import re          # Handle regular expressions for text cleanup and filtering
import time        # Measure elapsed time for performance tracking
import unicodedata
import subprocess  # Run external commands like `ebook-convert`
import shutil      # High-level file operations (copy, move, check binaries)
import argparse    # Parse command-line arguments provided by the user
import platform    # Detect the operating system for platform-specific handling
from PIL import Image, ImageFilter, ImageOps        # Image manipulation
from pdf2image import convert_from_path             # PDF to image conversion
from pytesseract import image_to_string, image_to_pdf_or_hocr    # Text recognition (OCR)
from tqdm import tqdm                                       # Progress bar
from docx import Document                                    # Word document manipulation
from reportlab.pdfgen import canvas                         # PDF generation
from reportlab.lib.pagesizes import A4                       # Standard page size
from reportlab.lib.units import cm                          # Measurement units
from PyPDF2       import PdfMerger                          # Merge multiple PDF byte streams into a single PDF document


# Two Tesseract configurations:
# - default: fast and accurate, does not preserve layout
# - layout-aware: tries to maintain original blocks and columns
# -l por: Portuguese language
tesseract_default_config = r'--oem 3 --psm 1 -l por'
tesseract_layout_config  = r'--oem 1 --psm 11 -l por'

def detect_package_manager():
    """Detects the system's package manager"""
    system = platform.system()
    if system == "Darwin":
        return "brew"
    elif system == "Linux":
        if shutil.which("apt"):
            return "apt"
        elif shutil.which("dnf"):
            return "dnf"
        elif shutil.which("yum"):
            return "yum"
    return None

def check_dependencies(generate_epub=False):
    """Checks if all system dependencies are installed and suggests install commands"""
    missing = []

    # Check essential system commands
    if shutil.which('tesseract') is None:
        missing.append('tesseract (OCR)')
    if shutil.which('pdftoppm') is None:
        missing.append('poppler-utils (pdftoppm for PDF to image conversion)')
    if generate_epub and shutil.which('ebook-convert') is None:
        missing.append('calibre (ebook-convert to create EPUB)')

    if missing:
        print("\n❌ Missing dependencies detected:")
        for item in missing:
            print(f"   - {item}")

        package_manager = detect_package_manager()

        print("\nℹ️ Install required packages using the following command:\n")

        if package_manager == "apt":
            print("   sudo apt install tesseract-ocr poppler-utils calibre\n")
        elif package_manager == "dnf":
            print("   sudo dnf install tesseract poppler-utils calibre\n")
        elif package_manager == "yum":
            print("   sudo yum install tesseract poppler-utils calibre\n")
        elif package_manager == "brew":
            print("   brew install tesseract poppler")
            print("   brew install --cask calibre\n")
            print("📌 If 'ebook-convert' is not found, add Calibre to your PATH:\n")
            print("   echo 'export PATH=\"$PATH:/Applications/calibre.app/Contents/MacOS\"' >> ~/.zshrc")
            print("   source ~/.zshrc\n")
        else:
            print("❓ Unrecognized system. Please install the missing dependencies manually.\n")

        exit(1)

def preprocess_image(img):
    """Pre-processes image to improve OCR quality"""
    img = img.convert('L')          # Convert to grayscale
    img = ImageOps.autocontrast(img)  # Auto contrast enhancement
    img = img.filter(ImageFilter.MedianFilter())  # Noise reduction filter
    return img

def clean_text_portuguese(text):
    """Cleans text by removing unwanted (non-ASCII) non-Portuguese special characters, preserving Portuguese common caracters"""
    allowed = (
        'a-zA-Z0-9'
        'áéíóúàãõâêôç'
        'ÁÉÍÓÚÀÃÕÂÊÔÇ'
        '\\s'
        '\\.,;:?!()\\[\\]{}\\-"\''  # basic punctuation
    )
    return re.sub(f'[^{allowed}]', '', text)

def extract_text_from_pdf(pdf_source_path, quiet=False):
    """Converts PDF to text using OCR"""
    pages = convert_from_path(pdf_source_path, dpi=400)  # Convert PDF to images (400 DPI)
    final_text = ""
    page_texts = []
    
    tesseract_custom_config=tesseract_default_config

    # Process each page with progress bar
    ocr_start = time.perf_counter()
    for page_img in tqdm(pages, desc="Extracting text (OCR)...", 
                        ascii=False, ncols=75, disable=quiet):
        page_img = preprocess_image(page_img)  # Pre-processing
        text = image_to_string(page_img, config=tesseract_custom_config)  # Extract text
        text = clean_text_portuguese(text) # clean non-ASCII preserving Portuguese caracters
        page_texts.append(text)
        final_text += text + "\n\n"
    
    ocr_time = time.perf_counter() - ocr_start
    return final_text, page_texts, ocr_time

def save_as_docx(text, output_path):
    """Saves extracted text to a DOCX file"""
    start = time.perf_counter()
    title = os.path.splitext(os.path.basename(output_path))[0].replace('_', ' ')
    document = Document()
    document.core_properties.title = title
    document.core_properties.author = "pdf2ocr"

    # Add each line as a paragraph in the document
    # for line in text.split('\n'):
    #     line = line.strip()
    #     if line:
    #         document.add_paragraph(line)

    #Preserve paragraphs
    paragraphs = text.strip().split('\n\n')
    for para in paragraphs:
        clean_para = para.strip().replace('\n', ' ')
        if clean_para:
            document.add_paragraph(clean_para)
    
    document.save(output_path)
    return time.perf_counter() - start

def remove_accents(texto):
    return unicodedata.normalize("NFKD", texto).encode("ASCII", "ignore").decode("ASCII")


def save_as_pdf(text_pages, output_path, filename):
    """Generates a new PDF with OCR-processed text"""
    start = time.perf_counter()
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    for idx, page_text in enumerate(text_pages, start=1):
        x = 2 * cm  # Left margin
        y = height - 3 * cm  # Top margin

        # Page header
        c.setFont("Helvetica-Bold", 10)
        header = f"OCR - {filename} - Pag. {idx}"
        header = remove_accents(header)
        c.drawCentredString(width / 2, height - 1 * cm, header)

        # Main text
        c.setFont("Helvetica", 10)
        line_height = 12

        #Preserve paragraphs
        paragraphs = page_text.strip().split('\n\n')
        for para in paragraphs:
            lines = para.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    c.drawString(x, y, line)
                    y -= line_height
                    if y < 2 * cm:  # Check page end
                        c.showPage()
                        y = height - 3 * cm
                        c.setFont("Helvetica", 10)
            y -= line_height  # Extra space between paragraphs

        c.showPage()

    c.save()
    return time.perf_counter() - start


def convert_docx_to_epub(docx_path, epub_path):
    """Converts DOCX to EPUB using Calibre"""
    start = time.perf_counter()
    try:
        title = os.path.splitext(os.path.basename(str(epub_path)))[0].replace('_', ' ')

        result = subprocess.run(
            [
                "ebook-convert", docx_path, epub_path,
                '--title', title,
                '--authors', 'pdf2ocr',
                '--comments', 'Converted by pdf2ocr',
                '--language', 'pt',
                '--level1-toc', "//h:h1"
            ],
            capture_output=True,
            text=True,
            check=True
        )
        return True, time.perf_counter() - start, result.stdout
    except subprocess.CalledProcessError as e:
        return False, time.perf_counter() - start, e.stderr

def process_pdfs_with_ocr(input_folder, output_folder, generate_docx, generate_pdf, 
                      generate_epub, generate_html, quiet=False, summary_output=False, log_path=None):
    """Main function that orchestrates the entire conversion process"""
    
    # Enable DOCX automatically if EPUB is requested
    if generate_epub and not generate_docx:
        if not quiet and not summary_output:
            print(f"\nℹ️ Note: EPUB requires DOCX as prerequisite. Enabling DOCX generation automatically.")
        generate_docx = True
    
    # Configure log file if specified
    log_file = None
    if log_path:
        log_file = open(log_path, "w", encoding="utf-8", buffering=1)
        log_file.write("\n=== Process Started ===\n\n")
        log_file.flush()

    # Create output folders as needed
    os.makedirs(output_folder, exist_ok=True)
    if generate_docx:
        docx_dir = os.path.join(output_folder, 'docx')
        os.makedirs(docx_dir, exist_ok=True)
        if log_file:
            log_file.write(f"✅ DOCX folder created - {docx_dir}/\n")
            log_file.flush()
    if generate_pdf:
        pdf_dir = os.path.join(output_folder, 'pdf_ocr')
        os.makedirs(pdf_dir, exist_ok=True)
        if log_file:
            log_file.write(f"✅ PDF folder created - {pdf_dir}/\n")
            log_file.flush()
    if generate_epub:
        epub_dir = os.path.join(output_folder, 'epub')
        os.makedirs(epub_dir, exist_ok=True)
        if log_file:
            log_file.write(f"✅ EPUB folder created - {epub_dir}/\n")
            log_file.flush()
    if generate_html:
        html_dir = os.path.join(output_folder, 'html')
        os.makedirs(html_dir, exist_ok=True)
        if log_file:
            log_file.write(f"✅ HTML folder created - {html_dir}/\n")
            log_file.flush()

    # List and sort PDF files in input folder
    pdf_files = [f for f in os.listdir(input_folder) if f.lower().endswith('.pdf')]
    pdf_files.sort()

    total = len(pdf_files)
    if total == 0:
        if not quiet:
            print("⚠️ No PDF files found!")
        return

    # Process each PDF file
    total_start = time.perf_counter()
    for idx, filename in enumerate(pdf_files, start=1):
        file_start = time.perf_counter()
        
        if not quiet and not summary_output:
            print(f"\n📄 Processing file {idx:02d} of {total:02d}: {filename}\n")

        if log_file:
            log_file.write(f"\n--- Processing file {idx:02d} of {total:02d}: {filename} ---\n")
            log_file.flush()

        try:
            pdf_source_path = os.path.join(input_folder, filename)
            base_name = os.path.splitext(filename)[0]
            
            # Step 1: Text extraction with OCR
            final_text, page_texts, ocr_time = extract_text_from_pdf(pdf_source_path, quiet or summary_output)
            
            if not quiet and not summary_output:
                print(f"\n🔹 Text extraction (OCR): completed in {ocr_time:.2f} seconds")
            
            if log_file:
                log_file.write(f"✅ OCR completed: {ocr_time:.2f}s\n")
                log_file.flush()

            # Step 2: Output file generation
            times = {}
            if not quiet and not summary_output:
                print("\n🔹 Creating output files:")

            # Generate DOCX if requested
            if generate_docx:
                docx_path = os.path.join(output_folder, docx_dir, base_name + '.docx')
                times['docx'] = save_as_docx(final_text, docx_path)
                if not quiet and not summary_output:
                    print(f"    - 📄 DOCX created in {times['docx']:.2f} seconds")
                if log_file:
                    log_file.write(f"✅ DOCX created: {times['docx']:.2f}s\n")
                    log_file.flush()

            # Generate PDF if requested
            if generate_pdf:
                pdf_path = os.path.join(output_folder, pdf_dir, base_name + '_ocr.pdf')
                times['pdf'] = save_as_pdf(page_texts, pdf_path, base_name)
                if not quiet and not summary_output:
                    print(f"    - 📄 OCR PDF created in {times['pdf']:.2f} seconds")
                if log_file:
                    log_file.write(f"✅ OCR PDF created: {times['pdf']:.2f}s\n")
                    log_file.flush()

            if generate_html:
                html_path = os.path.join(output_folder, html_dir, base_name + '.html')
                html_text = "".join(f"<p>{page.strip()}</p>\n" for page in page_texts)
                times['html'] = save_as_html(html_text, html_path)
                if log_file:
                    log_file.write(f"✅ HTML saved: {times['html']:.2f}s\n")
                    log_file.flush()

            # Generate EPUB if requested
            if generate_epub:
                epub_path = os.path.join(output_folder, epub_dir, base_name + '.epub')
                success, times['epub'], output = convert_docx_to_epub(docx_path, epub_path)
                
                if success:
                    if not quiet and not summary_output:
                        print(f"    - 📄 EPUB created in {times['epub']:.2f} seconds")
                    if log_file:
                        log_file.write(f"✅ EPUB created: {times['epub']:.2f}s\n")
                        log_file.write("\n=> ebook-convert output:\n\n")
                        log_file.write(output + "\n")
                        log_file.flush()
                else:
                    if not quiet and not summary_output:
                        print(f"    - ❌ Error creating EPUB ({times['epub']:.2f}s)")
                    if log_file:
                        log_file.write(f"❌ Error creating EPUB ({times['epub']:.2f}s): {output}\n")
                        log_file.flush()

            # Record total file processing time
            total_file_time = time.perf_counter() - file_start
            if not quiet and not summary_output:
                print(f"\n🔹 Total time for file: {total_file_time:.2f} seconds")

            if log_file:
                log_file.write(f"⏱️ Total file time: {total_file_time:.2f} seconds\n")
                #log_file.write(f"--- File completed: {filename} ---\n")
                log_file.write(f"--- File completed ---\n")
                log_file.flush()

        except Exception as e:
            # Error handling with time tracking
            total_file_time = time.perf_counter() - file_start
            if not quiet and not summary_output:
                print(f"\n❌ Error processing {filename}: {e}")
                print(f"🔹 Elapsed time: {total_file_time:.2f} seconds")
            if log_file:
                log_file.write(f"❌ Unexpected error: {e}\n")
                log_file.write(f"⏱️ Elapsed time: {total_file_time:.2f} seconds\n")
                log_file.write(f"--- File finished: {filename} ---\n\n")
                log_file.flush()

    # Final process summary
    total_end = time.perf_counter()
    total_duration = total_end - total_start
    
    if not quiet:
        print("\n🎯 Conversion summary:")
        print(f"📄 Total files found: {total}")
        print(f"⏱️ Total execution time: {total_duration:.2f} seconds")

    if log_file:
        log_file.write(f"\n=== Process completed ===\n\n")
        log_file.write(f"Total processed files: {total}\n")
        log_file.write(f"Total execution time: {total_duration:.2f} seconds\n")
        log_file.close()

def process_layout_pdf_only(source_dir, dest_dir,
                            tesseract_config,
                            quiet=False, summary=False, log_path=None):
    """
    Generate layout-preserving searchable PDFs for each PDF in source_dir,
    merging per-page PDF outputs via PyPDF2. Optimized for size by using lower DPI.
    """

    os.makedirs(dest_dir, exist_ok=True)
    pdf_dir = os.path.join(dest_dir, 'pdf_ocr_layout')
    os.makedirs(pdf_dir, exist_ok=True)

    log_file = None
    if log_path:
        log_file = open(log_path, "w", encoding="utf-8", buffering=1)
        log_file.write("=== Process Started ===\n\n")
        log_file.write(f"✅ Preserve-layout PDF folder created - {pdf_dir}\n\n")

    files = sorted(f for f in os.listdir(source_dir) if f.lower().endswith('.pdf'))
    total = len(files)
    if total == 0:
        if not quiet:
            print("⚠️ No PDF files found!")
        return

    overall_start = time.perf_counter()

    for idx, fname in enumerate(files, start=1):
        src_path = os.path.join(source_dir, fname)
        base = os.path.splitext(fname)[0]
        out_path = os.path.join(pdf_dir, f"{base}_ocr.pdf")

        print(f"\n📄 Processing file {idx:02d} of {total}: {fname}")
        if log_file:
            log_file.write(f"--- Processing file {idx:02d} of {total}: {fname} ---\n")

        # Use lower DPI to reduce image and PDF size
        pages = convert_from_path(src_path, dpi=200)
        print("Extracting text (OCR)...:", end=" ")
        ocr_start = time.perf_counter()
        pdf_pages = []

        for page_img in tqdm(pages, disable=quiet, ncols=75, leave=False):
            pdf_bytes = image_to_pdf_or_hocr(page_img, extension='pdf', config=tesseract_config)
            pdf_pages.append(pdf_bytes)

        ocr_time = time.perf_counter() - ocr_start
        print(f"\n\n🔹 Text extraction (OCR): completed in {ocr_time:.2f} seconds\n")
        if log_file:
            log_file.write(f"✅ OCR completed: {ocr_time:.2f}s\n")

        print("🔹 Creating output files:")
        merge_start = time.perf_counter()
        merger = PdfMerger()
        for b in pdf_pages:
            merger.append(io.BytesIO(b))
        with open(out_path, 'wb') as fout:
            merger.write(fout)
        pdf_time = time.perf_counter() - merge_start
        print(f"    - 📄 OCR PDF created in {pdf_time:.2f} seconds\n")
        if log_file:
            log_file.write(f"✅ PDF created: {pdf_time:.2f}s\n\n")

        file_total = ocr_time + pdf_time
        print(f"🔹 Total time for file: {file_total:.2f} seconds")
        if log_file:
            log_file.write("--- File completed ---\n\n")

    overall_time = time.perf_counter() - overall_start
    print(f"\n🎯 Conversion summary:")
    print(f"📄 Total files found: {total}")
    print(f"⏱️ Total execution time: {overall_time:.2f} seconds")

    if log_file:
        log_file.write("=== Process completed ===\n\n")
        log_file.write(f"Total processed files: {total}\n")
        log_file.write(f"Total execution time: {overall_time:.2f}s\n")
        log_file.close()



def save_as_html(text, output_path):
    """Saves extracted text to an HTML file"""
    start = time.perf_counter()
    try:
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>OCR Output</title>
</head>
<body>
<pre>
{text}
</pre>
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        elapsed = time.perf_counter() - start
        print(f"    - 📄 HTML saved: ({elapsed:.2f}s)")
    except Exception as e:
        print(f"❌ Failed to save HTML: {e}")

    return time.perf_counter() - start



def parse_arguments():
    """Configure and parse command line arguments"""
    parser = argparse.ArgumentParser(description="Python script to apply OCR on PDF files and generate output in DOCX, searchable PDF, HTML and EPUB formats.")
    parser.add_argument("source_dir", help="Source folder with PDF files")
    parser.add_argument("--dest-dir", help="Destination folder (optional, defaults to input folder)", default=None)
    parser.add_argument("--docx", action="store_true", help="Generate DOCX files")
    parser.add_argument("--pdf", action="store_true", help="Generate OCR-processed PDF files")
    parser.add_argument("--html", action="store_true", help="Generate html files")
    parser.add_argument("--epub", action="store_true", help="Generate EPUB files (auto-enables --docx if needed)")
    parser.add_argument("--preserve-layout", action="store_true", help="Preserve original document layout (PDF only)")
    parser.add_argument("--quiet", action="store_true", help="Run silently (no output)")
    parser.add_argument("--short-output", "--summary-output", dest="summary_output", action="store_true",
        help="Display only final conversion summary (short output mode)"
    )
    parser.add_argument("--logfile", help="Path to log file (optional)")
    return parser.parse_args()

def main():
    """Main entry point for the CLI"""
    args = parse_arguments()
    check_dependencies(generate_epub=args.epub)
    source_dir = os.path.expanduser(args.source_dir)
    dest_dir = os.path.expanduser(args.dest_dir) if args.dest_dir else source_dir

    # ─── FORCE PDF‐ONLY IF LAYOUT PRESERVING ───
    # Choose Tesseract config based on --preserve-layout
    if args.preserve_layout:
        if args.docx or args.epub or args.html:
            print("\n⚠️  Warning: --preserve-layout is only compatible with PDF output. Other formats have been disabled.")
            if args.logfile:
                with open(args.logfile, "a", encoding="utf-8") as log:
                    log.write("⚠️  Warning: --preserve-layout is only compatible with PDF output. Other formats have been disabled.\n\n")
        args.pdf, args.docx, args.epub, args.html = True, False, False, False
        tconfig = tesseract_layout_config
    else:
        tconfig = tesseract_default_config

    if not (args.docx or args.pdf or args.epub or args.html):
        print("⚠️ You must select at least one option: --docx, --pdf, --epub, --html")
        exit(1)
    
    # Layout‐preserving mode: only generate searchable PDFs via Tesseract CLI
    if args.preserve_layout:
        process_layout_pdf_only(
            source_dir,
            dest_dir,
            tesseract_config=tconfig,
            quiet=args.quiet,
            summary=args.summary_output,
            log_path=args.logfile
        )
    else: 
        process_pdfs_with_ocr(
            source_dir,
            dest_dir,
            args.docx,
            args.pdf,
            args.epub,
            args.html,
            quiet=args.quiet,
            summary_output=args.summary_output,
            log_path=args.logfile
        )

if __name__ == "__main__":
    main()

