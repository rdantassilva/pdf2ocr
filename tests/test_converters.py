import os
import pytest
from pdf2ocr.converters import save_as_docx, save_as_html
from docx import Document

def test_save_as_docx(tmp_path):
    """Test DOCX file generation and content"""
    output_file = tmp_path / "test.docx"
    test_text = """First paragraph
    with multiple lines.

    Second paragraph
    also with multiple lines.
    """
    
    # Save the document
    save_as_docx(test_text, str(output_file))
    
    # Verify file was created
    assert output_file.exists()
    
    # Read back and verify content
    doc = Document(output_file)
    # Normalize spaces in paragraphs
    paragraphs = [' '.join(p.text.split()) for p in doc.paragraphs if p.text.strip()]
    
    assert len(paragraphs) == 2
    assert "First paragraph with multiple lines" in paragraphs[0]
    assert "Second paragraph also with multiple lines" in paragraphs[1]
    
    # Verify document properties
    core_props = doc.core_properties
    assert core_props.author == "pdf2ocr"
    assert core_props.title == "test"

def test_save_as_docx_multiple_paragraphs(tmp_path):
    """Test DOCX file generation with multiple paragraphs"""
    output_file = tmp_path / "test.docx"
    test_text = """First paragraph
    with multiple lines.

    Second paragraph
    also with multiple lines.

    Third paragraph
    on multiple lines."""
    
    # Save the document
    save_as_docx(test_text, str(output_file))
    
    # Verify file was created
    assert output_file.exists()
    
    # Read back and verify content
    doc = Document(output_file)
    # Normalize spaces in paragraphs
    paragraphs = [' '.join(p.text.split()) for p in doc.paragraphs if p.text.strip()]
    
    assert len(paragraphs) == 3, "Expected three paragraphs separated by double newlines"
    assert "First paragraph with multiple lines" in paragraphs[0]
    assert "Second paragraph also with multiple lines" in paragraphs[1]
    assert "Third paragraph on multiple lines" in paragraphs[2]
    
    # Each paragraph should have a single run
    doc_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert len(doc_paragraphs[0].runs) == 1, "First paragraph should have one run"
    assert len(doc_paragraphs[1].runs) == 1, "Second paragraph should have one run"
    assert len(doc_paragraphs[2].runs) == 1, "Third paragraph should have one run"

def test_save_as_html(tmp_path):
    """Test HTML file generation and content"""
    output_file = tmp_path / "test.html"
    test_text = ["Line 1\nLine 2", "Paragraph 2"]  # List of pages
    
    # Save the HTML
    save_as_html(test_text, str(output_file))
    
    # Verify file was created
    assert output_file.exists()
    
    # Read back and verify content
    with open(output_file, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Check HTML structure
    assert "<!DOCTYPE html>" in content
    assert "<html lang=\"en\">" in content
    assert "<meta charset=\"UTF-8\">" in content
    assert "<div class=\"page\">" in content
    assert "<div class=\"page-header\">" in content
    assert "pdf2ocr - Page" in content
    
    # Check that paragraphs are inside page divs
    pages = content.split('<div class="page">')
    assert len(pages) > 1  # At least one page
    
    # Check content is present somewhere in the document
    assert "Line 1" in content
    assert "Line 2" in content
    assert "Paragraph 2" in content 